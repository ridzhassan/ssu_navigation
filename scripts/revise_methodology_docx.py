# -*- coding: utf-8 -*-
"""Revise Chapter II methodology & technical content to align with implemented system + SDLC."""
from pathlib import Path
import shutil

from docx import Document

DOC = Path(
    r"C:\Users\ridzh\OneDrive\Desktop\Undergrad Thesis\SSC Navigation"
    r"\6. Development of Campus Navigation with Augmented Reality for Sulu State College.docx"
)


def main() -> None:
    backup = DOC.with_suffix(".docx.bak")
    if not backup.exists():
        shutil.copy2(DOC, backup)
        print("Created backup:", backup)

    d = Document(str(DOC))

    # --- Chapter intro ---
    d.paragraphs[167].text = (
        "This chapter presents the methods and design processes used in developing the "
        "Augmented Reality Campus Navigation System for Sulu State College. It describes "
        "the technical background of the project, including the tools, frameworks, and "
        "technologies utilized during development. The chapter follows the System "
        "Development Life Cycle (SDLC): planning, analysis, design, implementation, testing, "
        "deployment, and evaluation. It explains the conceptual framework (IPO), the "
        "multi-tier system architecture (web client, administrator interface, REST API, and "
        "data store), software and hardware requirements, and the key features aligned with "
        "the implemented campus navigation platform. Together, these elements document how "
        "the system was engineered as an integrated information system rather than as a "
        "single standalone program."
    )

    # --- Technical Background ---
    d.paragraphs[170].text = (
        "The system is engineered as a client-server information system. The public-facing "
        "campus navigator and the administrator panel are implemented as React single-page "
        "applications using TypeScript and Vite. Maps are rendered with Leaflet using "
        "industry-standard basemaps (e.g., Esri satellite and OpenStreetMap). The user's "
        "position is obtained through the browser Geolocation API, which relies on the "
        "device's GNSS/GPS and, where available, other location signals."
    )
    d.paragraphs[171].text = (
        "The backend is a Node.js application built with Express. It exposes a RESTful API "
        "for points of interest (POIs), buildings, authentication for administrators, and "
        "navigation-related resources. Data persistence uses SQLite (embedded via sql.js "
        "in the reference implementation), which supports relational storage of POI "
        "coordinates, descriptions, and related entities without requiring a separate "
        "database server during development. For production, the API can be hosted on cloud "
        "platforms (e.g., Railway) while static front-end assets are deployed to web hosting."
    )
    d.paragraphs[172].text = (
        "Navigation logic on the client includes a campus path network (map matrix) "
        "represented as a graph; shortest-path computation (e.g., Dijkstra's algorithm) "
        "uses edge weights derived from geographic distances between nodes so that the "
        "displayed route follows walkable connections rather than a single straight line "
        "when the graph is connected."
    )
    d.paragraphs[173].text = (
        "Augmented reality (AR) remains part of the overall thesis vision: a complementary "
        "mobile AR client (for example, built with Unity, AR Foundation, and ARCore on "
        "Android) can consume the same coordinate and API data to overlay directional cues "
        "in the camera view. The web-based map and API therefore form the core system "
        "development deliverable that supplies authoritative campus data and routing inputs "
        "for both browser and AR clients."
    )
    d.paragraphs[174].text = (
        "Hardware support includes smartphones or tablets with a modern browser, network "
        "access for API calls, and GPS or equivalent location services enabled for live "
        "positioning. For future AR modules, devices must meet ARCore compatibility "
        "requirements (Android) or the equivalent platform AR stack."
    )
    d.paragraphs[175].text = (
        "GPS (or GNSS) provides latitude and longitude for the user, which the client "
        "compares against POI coordinates to compute distance, bearing, and turn-by-turn "
        "guidance text during active navigation."
    )
    d.paragraphs[176].text = (
        "Network connectivity is required for the web client to load map tiles and to "
        "retrieve or update data through the REST API unless a fully offline build is "
        "provided in a future iteration."
    )
    d.paragraphs[177].text = (
        "Device orientation sensors (where exposed to the browser) can support compass-style "
        "direction cues; dedicated IMU fusion is more prominent in native AR runtimes such "
        "as ARCore when implementing stable AR overlays."
    )
    d.paragraphs[178].text = (
        "Together, these technologies support interactive campus mapping, administrative "
        "data management, API-driven synchronization between Admin Panel and Client, "
        "graph-based outdoor routing on the map, and a path toward AR visualization that "
        "reuses the same backend and coordinate model."
    )

    # --- Methodology (SDLC / modified waterfall) ---
    d.paragraphs[185].text = (
        "The development of the Campus Navigation System followed the System Development "
        "Life Cycle (SDLC) in a structured, phase-by-phase manner. A modified waterfall model "
        "was applied: each phase produces documented outputs before the next begins, while "
        "short feedback loops during implementation allowed corrections to requirements and "
        "design without losing traceability to system objectives. This approach fits "
        "information-system projects where a stable data model, API contract, and "
        "deployment architecture must align before release."
    )

    d.paragraphs[187].text = (
        "This phase defines the problem (campus wayfinding at Sulu State College), scope, "
        "objectives, feasibility, and constraints. Deliverables include selection of the "
        "technology stack (React, Node/Express, SQLite, Leaflet), identification of actors "
        "(students, staff, administrators), and a work plan for analysis through deployment."
    )

    d.paragraphs[189].text = (
        "Functional requirements were derived for both administrator and end-user workflows: "
        "maintaining POIs and coordinates, searching destinations, displaying an interactive "
        "map, starting and stopping navigation, and computing distance and direction. "
        "Non-functional requirements include security for admin operations (e.g., "
        "authentication), CORS and deployment configuration for separated front-end and API "
        "hosts, response time expectations for API calls, and usability of the map interface "
        "on mobile browsers."
    )

    d.paragraphs[191].text = (
        "Design artifacts include the REST API structure, database schema for POIs and "
        "related entities, the campus path graph (nodes and edges) for routing, wireframes "
        "or screen flows for the Client and Admin Panel, and deployment topology (static "
        "hosting for front ends, Node hosting for the API). The design ensures a single "
        "source of truth for coordinates through the backend."
    )

    d.paragraphs[193].text = (
        "Implementation was carried out in layers: the Express API and persistence, the "
        "Admin Panel for data entry and verification, and the public Client with Leaflet "
        "map, markers, layer control, and navigation UI. Client-side utilities implement "
        "Haversine distance, bearing, and graph shortest-path routing. Optional AR client "
        "work can proceed in parallel once API and coordinates are stable."
    )

    d.paragraphs[195].text = (
        "Testing included manual and structured checks: API endpoints (e.g., POI retrieval "
        "and updates), synchronization of coordinate changes from Admin to Client, map tile "
        "loading and layout, geolocation behavior, navigation state transitions, and "
        "regression after fixes. Documented test cases support traceability to requirements."
    )

    d.paragraphs[197].text = (
        "Deployment consists of building production bundles for the Client and Admin Panel, "
        "hosting static files on a web server, running the Node API on a process manager or "
        "cloud service (e.g., Railway), configuring environment variables and CORS origins, "
        "and enabling HTTPS. End users access the navigator via URL; administrators use "
        "the secured Admin Panel."
    )

    d.paragraphs[199].text = (
        "Evaluation compares system behavior against objectives, gathers user or stakeholder "
        "feedback, and records limitations (e.g., GPS accuracy, outdoor-only graph, network "
        "dependency). Results inform maintenance, data cleanup, and future enhancements such "
        "as indoor routing or full AR integration."
    )

    # SDLC phase headings (systems development terminology)
    d.paragraphs[186].text = "1. Planning"
    d.paragraphs[188].text = "2. Systems Analysis and Requirements Determination"
    d.paragraphs[190].text = "3. System Design"
    d.paragraphs[192].text = "4. Implementation (Coding and Integration)"
    d.paragraphs[194].text = "5. Testing and Quality Assurance"
    d.paragraphs[196].text = "6. Deployment and Installation"
    d.paragraphs[198].text = "7. Evaluation and Maintenance"

    # --- IPO ---
    d.paragraphs[207].text = (
        "The system receives administrator-defined and API-served data (POI names, types, "
        "descriptions, latitude/longitude), the user's real-time geographic position from "
        "the Geolocation API, map tile requests for the selected basemap, and user actions "
        "(search, select destination, start/stop navigation). For a future AR client, "
        "inputs also include the live camera stream and device pose."
    )

    d.paragraphs[210].text = (
        "REST API consumption: the client requests campus data and receives JSON payloads "
        "for POIs and related resources."
    )
    d.paragraphs[211].text = (
        "Map rendering: coordinates are projected onto the basemap; markers and polylines "
        "show locations and computed routes."
    )
    d.paragraphs[212].text = (
        "Path computation: the campus graph is queried for a shortest path between the "
        "nearest graph nodes to the user and destination, producing a sequence of waypoints "
        "for display; fallback straight-line routing applies if the graph is unavailable."
    )
    d.paragraphs[213].text = (
        "Presentation logic: search filtering, selection state, navigation instructions "
        "(distance, direction text), and administrator CRUD workflows."
    )

    d.paragraphs[215].text = (
        "Outputs include the interactive map with POI markers, highlighted routes, "
        "navigation guidance, and synchronized updates after administrative edits. The "
        "Admin Panel outputs updated records persisted in the database. An AR module would "
        "additionally output registered overlays (e.g., arrows) aligned to the real "
        "environment."
    )

    # --- System architecture (heading + body pairs) ---
    d.paragraphs[221].text = "Web Presentation Layer"
    d.paragraphs[222].text = (
        "This layer comprises the public React Client (campus map, search, navigation, "
        "layer switcher) and the React Admin Panel (authentication, POI and campus data "
        "management). It is delivered as static web assets and runs in the browser."
    )

    d.paragraphs[223].text = "Augmented Reality Layer (Complementary Extension)"
    d.paragraphs[224].text = (
        "This layer corresponds to a native or Unity-based mobile client using ARCore and "
        "AR Foundation. It processes the camera feed and device motion to anchor virtual "
        "arrows or labels. It relies on the same POI coordinates and optional route data "
        "provided by the backend API as the web client."
    )

    d.paragraphs[225].text = "Application Layer (REST API)"
    d.paragraphs[226].text = (
        "The Node.js Express server implements REST endpoints, validation, authentication "
        "middleware, and error handling. It is the single integration point for both web "
        "front ends and any future AR client."
    )

    d.paragraphs[227].text = "Navigation Services (Client-Side)"
    d.paragraphs[228].text = (
        "Geolocation tracking, Haversine distance, bearing, turn-by-turn text, and "
        "graph-based shortest-path routing over the campus walkability model execute in the "
        "browser for the web client, using POI positions from the API."
    )

    d.paragraphs[229].text = "Data Persistence Layer"
    d.paragraphs[230].text = (
        "SQLite stores POIs, buildings, administrative users, route definitions, and "
        "related metadata. The database is accessed through the API so that all clients "
        "share one authoritative dataset."
    )

    d.paragraphs[232].text = (
        "This section outlines the software and hardware requirements for developing and "
        "running the system: modern web stack tools for the Client and Admin Panel, Node.js "
        "for the API, SQLite for persistence, and devices with GPS and a compatible browser "
        "(plus ARCore-capable hardware when the AR extension is implemented)."
    )

    # --- Software table (first table in document: header + 5 data rows) ---
    tbl = d.tables[0]
    rows = [
        ("Software", "Description"),
        (
            "Node.js (LTS)",
            "JavaScript runtime used to execute the Express REST API server in development and production.",
        ),
        (
            "Express.js",
            "Web application framework for routing, middleware (e.g., CORS, security headers), and JSON APIs.",
        ),
        (
            "SQLite / sql.js",
            "Relational storage for POIs, buildings, admins, and routes; embedded engine suitable for portable deployment.",
        ),
        (
            "React, TypeScript, Vite, Leaflet",
            "Front-end stack and map libraries for the Client and Admin Panel (markers, routes, basemap layers).",
        ),
        (
            "Modern browser + Geolocation API",
            "End-user runtime for the web navigator; provides GPS/GNSS position and orientation where available.",
        ),
    ]
    for i, (a, b) in enumerate(rows):
        if i < len(tbl.rows):
            tbl.rows[i].cells[0].text = a
            tbl.rows[i].cells[1].text = b

    # --- Implementation feature list ---
    d.paragraphs[241].text = "REST API and Single Source of Truth"
    d.paragraphs[242].text = (
        "The Express API exposes POI, building, authentication, and route resources so that "
        "the Admin Panel and Client always read and write consistent coordinates and "
        "descriptions."
    )

    d.paragraphs[243].text = "Real-Time Browser Geolocation"
    d.paragraphs[244].text = (
        "During navigation, the Client uses the Geolocation API to track the user, "
        "recalculate distance and bearing, and update on-screen guidance."
    )

    d.paragraphs[245].text = "Interactive Campus Map (Leaflet)"
    d.paragraphs[246].text = (
        "Basemap layers (e.g., Esri aerial, street maps), POI markers, optional campus path "
        "network overlay, and route polylines are rendered on the map."
    )

    d.paragraphs[247].text = "Search, Selection, and Navigation UI"
    d.paragraphs[248].text = (
        "Users search POIs, select a destination, start or stop navigation, and view "
        "distance and estimated walking time in a mobile-friendly panel."
    )

    d.paragraphs[249].text = "Graph-Based Outdoor Routing (Map Matrix)"
    d.paragraphs[250].text = (
        "A campus walkability graph connects intersections and anchors; Dijkstra's algorithm "
        "computes shortest paths between nearest nodes so routes follow defined walkways "
        "when the graph is connected."
    )

    d.save(str(DOC))
    print("Updated:", DOC)


if __name__ == "__main__":
    main()
